#txt_to_docx.py

 	

from docx import Document
from docx.shared import Inches

document = Document()

f = open('txt/input.txt', 'r', encoding='utf-8')
# str = unicode(f, "utf-8")
p = document.add_paragraph(f)
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True
f.close()

document.add_page_break()
document.save('demo.docx')