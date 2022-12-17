# url_doc.py

import urllib
import requests
from bs4 import BeautifulSoup

def download_url():
	url = input()
	resp = requests.get(url)
	html = resp.text
	soup = BeautifulSoup(html, "html.parser")
	tags = soup.find_all(['h2', 'p'])
	f = open('some4.html', 'w', encoding='utf-8')
	f.write(soup.prettify())
	f.close()
	"""f = open('tags.html', 'w')
	f.write(tags.prettify())
	f.close()"""
	print("URL скачан!")
	find_p()


import json

def find_p():
	f = open('some4.html', 'r')
	soup = BeautifulSoup(f, features='lxml')
	divs = soup.find_all('div')
	with open('txt/input.txt', 'w', encoding='utf-8') as fw:
		for div in divs:
		    ps = div.find_all('p')
		    for p in ps:
		    	fw.write(p.get_text("\n"))
	print('Теги "p" найдены')
	to_docx()

from docx import Document
from docx.shared import Inches

def to_docx():
	document = Document()

	f = open('txt/input.txt', 'r', encoding='utf-8')
	# str = unicode(f, "utf-8")
	p = document.add_paragraph(f)
	p.add_run('bold').bold = True
	p.add_run(' and some ')
	p.add_run('italic.').italic = True
	f.close()

	document.add_page_break()
	document.save('demo.docx')
	print("Docx создан! Готово!")

download_url()