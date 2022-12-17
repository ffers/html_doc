#txt_docx.py

from docx import Document
import re
import os

path = '/home/den/prog/html_doc/txt/'
direct = os.listdir(path)

for i in direct:
    document = Document()
    document.add_heading(i, 0)
    myfile = open('/home/den/prog/html_doc/txt/'+i).read()
    myfile = re.sub(r'[^\x00-\x7F]+|\x0c',' ', myfile) # remove all non-XML-compatible characters
    p = document.add_paragraph(myfile)
    document.save('/path/to/write/to/'+i+'.docx')